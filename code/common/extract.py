from io import StringIO

from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from docx import Document

def extract_pdf(pdf_file = r'C:\Users\Namita\OneDrive\Desktop\major_project\code\data\Teacher Resume Samples .pdf'):
    output_string = StringIO()
    with open(pdf_file, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

    return output_string.getvalue()

def extract_doc(doc_file=""):
    if not doc_file:
        print("Error: Please provide the path to a Word document.")
        return

    try:
        # Open the Word document
        doc = Document(doc_file)
        text = ''

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'

        return text

    except Exception as e:
        print(f"Error: {e}")
        return None

# Example usage
if __name__ == "__main__":
    docx_file_path = r'C:\Users\Namita\OneDrive\Desktop\major_project\code\data\report minor project.docx'
    text_content = extract_doc(doc_file=docx_file_path)

    if text_content is not None:
        print("DOC Extraction Result:")
        print(text_content)
    else:
        print("DOC Extraction failed.")